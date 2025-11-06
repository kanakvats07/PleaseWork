import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
import io
import datetime
import json
import re

# ---------------------------
# CONFIGURE - Replace API KEY
# ---------------------------
API_KEY = "AIzaSyAOdr8vsKZSmloEyZKS2U1Lb8cGK2iXwno"  # <-- replace this with your actual Google API key
genai.configure(api_key=API_KEY)

# ---------------------------
# UI: Light Ocean Background Theme
# ---------------------------
BACKGROUND_IMAGE = "https://images.unsplash.com/photo-1507525428034-b723cf961d3e"  # Light ocean wave image
st.markdown(
    f"""
    <style>
    .stApp {{
      background-image: url("{BACKGROUND_IMAGE}");
      background-size: cover;
      background-attachment: fixed;
      background-position: center;
      color: #043A47 !important;  /* Deep teal for visibility */
    }}
    .content-container {{
      background: rgba(255, 255, 255, 0.8); /* Light transparent box */
      padding: 24px;
      border-radius: 12px;
      color: #043A47;  /* Match text with background */
    }}
    .stTextInput, .stTextArea, .stButton > button {{
      color: #043A47 !important;
    }}
    h1, h2, h3, h4, h5, h6 {{
      color: #022E38 !important; /* Darker ocean shade */
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

st.set_page_config(page_title="Transcript → Professional BRD", page_icon="📄", layout="wide")

st.markdown('<div class="content-container">', unsafe_allow_html=True)
st.title("📝 Transcript → Professional BRD Generator")
st.write("Paste your meeting transcript. The app will ask Gemini for a structured BRD JSON and build a clean Word (.docx) with native tables and an ASCII architecture diagram.")

transcript = st.text_area("Paste meeting transcript here:", height=360, placeholder="Paste Teams transcript...")

model_choice = st.selectbox("Model", options=["models/gemini-2.5-pro"], index=0)
doc_name = st.text_input("Word filename", value="Business_Requirements_Document.docx")


# ---------------------------
# Helpers: JSON extraction
# ---------------------------
def strip_code_fences(text: str) -> str:
    """Remove ``` ``` or ```json ``` fences if present."""
    # remove triple backtick fences
    text = re.sub(r"```(?:json)?\s*", "", text)
    text = text.replace("```", "")
    return text.strip()

def find_json_substring(text: str) -> str:
    """
    Find the first top-level JSON object in text by locating the first '{' and
    finding its matching closing '}'. Returns substring or raises ValueError.
    """
    start = text.find("{")
    if start == -1:
        raise ValueError("No opening brace found for JSON.")
    depth = 0
    end = -1
    for i in range(start, len(text)):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i
                break
    if end == -1:
        raise ValueError("Matching closing brace for JSON not found.")
    return text[start:end+1]

def extract_json(text: str):
    """
    Try multiple strategies to extract JSON from text:
    1) remove code fences and try json.loads(text)
    2) find JSON substring using find_json_substring and json.loads
    3) try to replace single quotes with double quotes as a last resort
    Returns Python object or raises Exception.
    """
    if not text or not isinstance(text, str):
        raise ValueError("Empty or invalid text for JSON extraction.")
    t = strip_code_fences(text)
    # First attempt: direct load
    try:
        return json.loads(t)
    except Exception:
        pass
    # Second: locate JSON substring by matching braces
    try:
        js_sub = find_json_substring(t)
        return json.loads(js_sub)
    except Exception:
        pass
    # Third: naive replace single quotes -> double quotes (risky), then try to load
    try:
        t2 = t.replace("'", '"')
        return json.loads(t2)
    except Exception as e:
        raise ValueError(f"Failed to extract JSON from model output. Last error: {e}")


# ---------------------------
# Call Gemini - structured JSON
# ---------------------------
def call_gemini_structured(transcript: str, model_name: str = "models/gemini-2.5-pro"):
    """
    Ask Gemini to return a JSON structure containing the BRD sections and an ascii_diagram.
    We use a strict instruction but also handle cases the model returns fences or commentary.
    """
    prompt = f"""
You are an expert Business Analyst. Given the following meeting TRANSCRIPT, produce ONLY valid JSON (no commentary). The JSON MUST have these keys:

- executive_summary: (string)
- objectives: (string)
- scope_in: (array of strings)
- scope_out: (array of strings)
- stakeholders: (array of objects with fields: name, title, role)
- functional_requirements: (array of objects with fields: id, description, priority)
- nonfunctional_requirements: (array of strings)
- technical_architecture: (string)
- timeline: (string)
- risks: (array of objects with fields: id, description, mitigation)
- ascii_diagram: (string)  # Provide an ASCII diagram with boxes and arrows showing data flow

Transcript:
\"\"\"{transcript}\"\"\"


Important constraints:
1) Return valid JSON only. Do NOT add any markdown, explanation, or text before/after the JSON.
2) Keep strings concise but informative.
3) For ascii_diagram, use simple monospace box/arrow style:
   Example: [CRM] --> [Ingest] --> [Lakehouse] --> [Power BI]
Return the JSON now.
"""
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)
    return response.text


# ---------------------------
# Fallback: ask for ASCII diagram cleanly
# ---------------------------
def call_gemini_ascii(technical_text: str, model_name: str = "models/gemini-2.5-flash"):
    """
    If the primary JSON lacked ascii_diagram, ask Gemini to produce ONLY an ASCII diagram.
    """
    prompt = f"""
You are an expert solution architect. Given this short technical summary, produce ONLY a small ASCII architecture diagram (no text explanation, no JSON) using boxes and arrows, monospace-friendly.

Technical summary:
\"\"\"{technical_text}\"\"\"


Output example:
[CRM] --> [Ingest] --> [Lakehouse] --> [Semantic Model] --> [Power BI]

Now output ONLY the ASCII diagram.
"""
    model = genai.GenerativeModel(model_name)
    resp = model.generate_content(prompt)
    return resp.text.strip()


# ---------------------------
# Build Word document from parsed JSON
# ---------------------------
def build_word_doc_from_json(brd_json: dict, transcript_text: str) -> Document:
    doc = Document()

    # Title and meta
    doc.add_heading("BUSINESS REQUIREMENTS DOCUMENT", 0)
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Source Transcript Length: {len(transcript_text)} characters")
    doc.add_paragraph("")

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(brd_json.get("executive_summary", "N/A"))

    # Objectives
    doc.add_heading("2. Project Objectives", level=1)
    doc.add_paragraph(brd_json.get("objectives", "N/A"))

    # Scope
    doc.add_heading("3. Scope Definition", level=1)
    doc.add_heading("3.1 In-Scope", level=2)
    scope_in = brd_json.get("scope_in", [])
    if scope_in:
        for s in scope_in:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s)
    else:
        doc.add_paragraph("N/A")

    doc.add_heading("3.2 Out-of-Scope", level=2)
    scope_out = brd_json.get("scope_out", [])
    if scope_out:
        for s in scope_out:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s)
    else:
        doc.add_paragraph("N/A")

    # Stakeholders table
    doc.add_heading("4. Stakeholder Analysis", level=1)
    stakeholders = brd_json.get("stakeholders", [])
    if stakeholders:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Name"
        hdr[1].text = "Title"
        hdr[2].text = "Role in Project"
        for s in stakeholders:
            row = table.add_row().cells
            row[0].text = s.get("name", "")
            row[1].text = s.get("title", "")
            row[2].text = s.get("role", "")
    else:
        doc.add_paragraph("No stakeholders identified.")

    # Functional Requirements table
    doc.add_heading("5. Functional Requirements", level=1)
    frs = brd_json.get("functional_requirements", [])
    if frs:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Req ID"
        hdr[1].text = "Description"
        hdr[2].text = "Priority"
        for fr in frs:
            row = table.add_row().cells
            row[0].text = fr.get("id", "")
            row[1].text = fr.get("description", "")
            row[2].text = fr.get("priority", "")
    else:
        doc.add_paragraph("No functional requirements captured.")

    # Non-functional requirements
    doc.add_heading("6. Non-Functional Requirements", level=1)
    nfs = brd_json.get("nonfunctional_requirements", [])
    if nfs:
        for nf in nfs:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(nf)
    else:
        doc.add_paragraph("N/A")

    # Technical architecture
    doc.add_heading("7. Technical Architecture", level=1)
    doc.add_paragraph(brd_json.get("technical_architecture", "N/A"))

    # ASCII diagram page
    doc.add_page_break()
    doc.add_heading("Architecture Diagram (ASCII)", level=1)
    ascii_diagram = brd_json.get("ascii_diagram", "").strip()
    if ascii_diagram:
        p = doc.add_paragraph()
        run = p.add_run(ascii_diagram)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    else:
        doc.add_paragraph("[ASCII diagram not generated]")

    # Timeline
    doc.add_heading("8. Timeline & Milestones", level=1)
    doc.add_paragraph(brd_json.get("timeline", "N/A"))

    # Risks table
    doc.add_heading("9. Risks & Mitigations", level=1)
    risks = brd_json.get("risks", [])
    if risks:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Risk ID"
        hdr[1].text = "Description"
        hdr[2].text = "Mitigation"
        for r in risks:
            row = table.add_row().cells
            row[0].text = r.get("id", "")
            row[1].text = r.get("description", "")
            row[2].text = r.get("mitigation", "")
    else:
        doc.add_paragraph("No risks captured.")

    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix: Transcript Excerpt", level=1)
    excerpt = transcript_text[:1200] + ("..." if len(transcript_text) > 1200 else "")
    p = doc.add_paragraph(excerpt)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(10)

    return doc


# ---------------------------
# Main button action
# ---------------------------
if st.button("🚀 Generate BRD (clean Word + ASCII diagram)"):
    if not transcript.strip():
        st.error("Please paste a transcript first.")
    else:
        with st.spinner("Requesting structured BRD from Gemini..."):
            try:
                raw_response = call_gemini_structured(transcript, model_choice)
            except Exception as e:
                st.error(f"Model call failed: {e}")
                raw_response = None

        if not raw_response:
            st.error("No response from model.")
        else:
            st.subheader("Raw model output (for debugging)")
            st.code(raw_response[:5000], language="")

            # Attempt to extract JSON robustly
            parsed = None
            try:
                parsed = extract_json(raw_response)
                st.success("✅ Parsed JSON from model.")
                st.subheader("Parsed JSON preview")
                st.json(parsed)
            except Exception as e:
                st.warning(f"Could not parse JSON from model output: {e}")
                # fallback: wrap raw into executive_summary
                parsed = {
                    "executive_summary": raw_response,
                    "objectives": "",
                    "scope_in": [],
                    "scope_out": [],
                    "stakeholders": [],
                    "functional_requirements": [],
                    "nonfunctional_requirements": [],
                    "technical_architecture": "",
                    "timeline": "",
                    "risks": [],
                    "ascii_diagram": ""
                }

            # If ascii_diagram missing or empty, attempt a focused call to create it
            if not parsed.get("ascii_diagram"):
                st.info("ASCII diagram missing — requesting a focused ascii diagram from the model...")
                try:
                    tech_summary = parsed.get("technical_architecture", "") or transcript[:1000]
                    ascii_only = call_gemini_ascii(tech_summary)
                    # Clean ascii_only (strip code fences)
                    ascii_only = strip_code_fences(ascii_only)
                    if ascii_only:
                        parsed["ascii_diagram"] = ascii_only
                        st.success("ASCII diagram generated via focused prompt.")
                        st.text(ascii_only)
                    else:
                        st.warning("Focused ASCII call returned empty.")
                except Exception as e:
                    st.warning(f"Failed to generate ascii diagram: {e}")

            # Ensure keys exist and are of proper types (defensive)
            def ensure_list(k):
                v = parsed.get(k, [])
                return v if isinstance(v, list) else []
            parsed["scope_in"] = ensure_list("scope_in")
            parsed["scope_out"] = ensure_list("scope_out")
            parsed["stakeholders"] = ensure_list("stakeholders")
            parsed["functional_requirements"] = ensure_list("functional_requirements")
            parsed["nonfunctional_requirements"] = ensure_list("nonfunctional_requirements")
            parsed["risks"] = ensure_list("risks")

            # Build Word doc
            doc = build_word_doc_from_json(parsed, transcript)
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            st.success("✅ BRD generated. Download the Word file below.")
            st.download_button(
                "📥 Download BRD (Word)",
                data=doc_bytes.getvalue(),
                file_name=doc_name or "Business_Requirements_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown('</div>', unsafe_allow_html=True)
